"""Resume parsing service - extracts text from PDF and DOCX files."""

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text content from a PDF file."""
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text content from a DOCX file."""
    doc = Document(file_path)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)


def extract_text(file_path: Path) -> str:
    """Extract text from a resume file (PDF or DOCX)."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif suffix == ".txt":
        return file_path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Use PDF, DOCX, or TXT.")


def extract_email(text: str) -> str | None:
    """Extract email address from resume text."""
    pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> str | None:
    """Extract phone number from resume text."""
    patterns = [
        r"\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
        r"\+?\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{0,4}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return None


def extract_name(text: str) -> str:
    """Extract candidate name from resume text (heuristic: first non-empty line)."""
    lines = text.strip().split("\n")
    for line in lines:
        cleaned = line.strip()
        if cleaned and not re.match(r"^[a-zA-Z0-9._%+-]+@", cleaned) and len(cleaned) < 60:
            # Skip lines that look like headers/titles
            if not any(
                kw in cleaned.lower()
                for kw in ["resume", "curriculum", "cv", "objective", "summary"]
            ):
                return cleaned
    return "Unknown Candidate"


def extract_skills(text: str) -> str | None:
    """Extract skills section from resume text."""
    skill_headers = [
        r"(?i)(?:technical\s+)?skills?",
        r"(?i)competenc(?:ies|e)",
        r"(?i)technologies",
        r"(?i)proficienc(?:ies|y)",
    ]
    text_lower = text.lower()
    for header_pattern in skill_headers:
        match = re.search(header_pattern, text)
        if match:
            start = match.end()
            # Find the next section header
            next_section = re.search(
                r"\n\s*(?:experience|education|projects|certif|work|employment|"
                r"objective|summary|references|awards|publications)",
                text_lower[start:],
            )
            end = start + next_section.start() if next_section else min(start + 500, len(text))
            skills_text = text[start:end].strip()
            if skills_text:
                return skills_text
    return None


def extract_experience_years(text: str) -> float | None:
    """Estimate years of experience from resume text."""
    # Look for explicit mentions
    patterns = [
        r"(\d+)\+?\s*years?\s*(?:of\s+)?(?:experience|exp)",
        r"(?:experience|exp)\s*:?\s*(\d+)\+?\s*years?",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return float(match.group(1))

    # Count date ranges in experience section
    date_pattern = r"(20\d{2}|19\d{2})\s*[-–]\s*(20\d{2}|19\d{2}|present|current)"
    year_ranges = re.findall(date_pattern, text, re.IGNORECASE)
    if year_ranges:
        total_years = 0.0
        for start_year, end_year in year_ranges:
            start = int(start_year)
            end = 2024 if end_year.lower() in ("present", "current") else int(end_year)
            total_years += max(0, end - start)
        return total_years if total_years > 0 else None

    return None


def extract_education(text: str) -> str | None:
    """Extract education section from resume text."""
    edu_match = re.search(r"(?i)education", text)
    if edu_match:
        start = edu_match.end()
        text_lower = text.lower()
        next_section = re.search(
            r"\n\s*(?:experience|skills|projects|certif|work|employment|"
            r"objective|summary|references|awards)",
            text_lower[start:],
        )
        end = start + next_section.start() if next_section else min(start + 500, len(text))
        edu_text = text[start:end].strip()
        if edu_text:
            return edu_text
    return None


def parse_resume(file_path: Path) -> dict:
    """Parse a resume file and extract structured information."""
    text = extract_text(file_path)
    return {
        "text": text,
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills(text),
        "experience_years": extract_experience_years(text),
        "education": extract_education(text),
    }
