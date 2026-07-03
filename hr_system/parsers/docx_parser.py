"""DOCX resume parser using python-docx."""

from pathlib import Path

from docx import Document


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text content from a DOCX file."""
    doc = Document(str(file_path))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Also extract from tables (common in resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError(f"Could not extract text from DOCX: {file_path.name}")
    return text
